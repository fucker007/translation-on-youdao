#coding:utf-8
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import *
import re
import time
import sys
def get_driver():
    driver = webdriver.Chrome()
    driver.get("http://www.youdao.com/w/eng/is/#keyfrom=dict2.index")
    driver.maximize_window()
    return driver
#查询单词
def search_word(driver,word):
    driver.delete_all_cookies()
    wait = WebDriverWait(driver, 60)
    element = wait.until(
        EC.presence_of_element_located((By.ID, "query"))
    )
    element.clear()
    element.send_keys(word)
    click_b = wait.until(
        EC.presence_of_element_located((By.ID, "f"))
    )
    try:
        sumbit = wait.until(
            #EC.element_to_be_clickable((By.XPATH,"//form[@id='f']/input[@type='submit']"))
            EC.element_to_be_clickable((By.CLASS_NAME,'s-btn'))
        )
        sumbit.submit()
    except ElementClickInterceptedException:
        search_word(driver,word)
    return driver


    #sumbit = click_b.find_element_by_xpath("./input[@type='submit']")
# 单词的意思
def get_meas(driver):
    wait = WebDriverWait(driver,20)
    divs = wait.until(
        EC.presence_of_element_located((By.ID, "results-contents"))
    )
    urs = divs.find_element_by_class_name('trans-container').find_element_by_xpath('./ul')
    #list.append("单词意思")
    #list.append(urs.text)
    mean_all = urs.text
    mean_all.encode('utf-8')
    return mean_all

# 找出单词的词根，并写入文件
def get_root_word(driver):
    time.sleep(1)
    root_word_click = driver.find_element_by_link_text('同根词')
    root_word_click.click()  # 点击同根词开关
    print(root_word_click.text)
    # 获取同根词的相应内容
    relWordTab = driver.find_element_by_id("relWordTab")
    relWordTab = relWordTab.text
    relWordTab.encode('utf-8')
    print(relWordTab)
    return relWordTab


# 词组短语，写入文件
def get_word_group(driver):
    wordGroup_click = driver.find_element_by_link_text('词组短语')
    wordGroup_click.click()
    wordGroup = driver.find_element_by_id("wordGroup").text
    wordGroup.encode('utf-8')
    return wordGroup

# 词语辨析的相关东西
def get_word_context():
    list = []
    discriminate_click = driver.find_element_by_link_text('词语辨析')
    discriminate_click.click()
    discriminate = driver.find_element_by_id("discriminate")
    context = discriminate.find_elements_by_class_name("wt-container")
    print ("词语辨析")
    for some in context:
        list.append(some.text)
        print (some.text)
    return list
        # context_title = context.find_elements_by_class_name("title")
        # context_word = context.find_elements_by_class_name("collapse-content")

# 用词造句，截取专业释意的句子，写入文件
def get_authority(driver):
    list = []
    authority_click = driver.find_element_by_link_text('权威例句')
    authority_click.click()
    authority = driver.find_element_by_id("authority").find_element_by_xpath("./ul")
    first = authority.find_elements_by_xpath("./li[1]/p[1]")
    second = authority.find_element_by_xpath("./li[2]/p[1]")
    print ("权威例句")
    print(first[0].text, second.text)
    list.append(first[0].text)
    list.append(second.text)
    return list

def get_word_from_file(file):
    list = []
    with open(file) as  f1:  # 打开'weibo_train_data.txt'文件
        f11 = f1.readlines()  # 将打开文件的内容读到内存中，with 在执行完命令后，会关闭文件
        for x in f11:
            x.strip()  # 除去每行的换行符
            if len(x.split(' ')) > 2:
                # x = x.split(' ')#文本分割，以table键分割
                # x = re.compile(r'([a-z)+')
                x = re.search(r'[a-z]+', x)
                x = x.group(0)
                # 将x的值存到一个数组当中，然后对这个数组进行查询
                #           print(x.group(0))
                list.append(x)
                print(x)
    return list


#search_word('urban')
driver = get_driver()
list1 = get_word_from_file("../wordfile/145_160.txt")
list2 = []
import docx
from docx import Document
from docx.shared import Inches
document = Document()
document.add_heading('高中高频词汇 ',0)
#for word in list1:
#    search_word(driver,word)
means_data = ""
for word in list1:
    document.add_heading(word,level=1)
    driver = search_word(driver,word)
    #means = []
    means = get_meas(driver)
    '''
    while means_data == means:
        search_word(driver,word)
        means = get_meas(driver)
    means_data = means
    '''
    print('*************************************************************')
    print ('单词的意思：', means)
    list2.append("单词的意思\n")
    list2.append(means)
    document.add_paragraph(means,style='IntenseQuote')
    #获取单词词根
    try:
        #root_word = []
        root_word = get_root_word(driver)
        print('单词的词根', root_word)
        list2.append("词根\n")
        list2.append(root_word)
        document.add_paragraph(
            root_word, style='ListBullet'
        )
    except NoSuchElementException:
        pass
    #获取单词短语
    try:

        #word_group = []
        word_group = get_word_group(driver)
        print ('单词的短语', word_group)
        list2.append("短语\n")
        list2.append(word_group)
        document.add_paragraph(
            word_group, style='ListBullet'
        )
    except NoSuchElementException:
        pass
    #获取详解
    try:
       # list_context = []
        list_context = get_word_context()
        print (list_context)
        list2.append("详细分析\n")
        for e in list_context:
            list2.append(e)
            document.add_paragraph(
                e, style='ListBullet'
            )
    except NoSuchElementException:
        pass
    #获取造句
    try:
        #list_authority = []
        list_authority  = get_authority(driver)
        list2.append("造句\n")
        for x in list_authority:
            list2.append(x)
            document.add_paragraph(
                '造句', style='ListBullet'
            )
            document.add_paragraph(
                x, style='ListBullet'
            )
        print ('造句', list_authority)
    except NoSuchElementException:
        pass
    print ('****************************************************************')
    continue
document.save('../wordfile/list29.docx')
print ("00000000000000000000000000000000000000000000000000000000000")
print(list2)
    #means.encode("utf-8").decode("gbk")




'''
for word in list1:
    list2 = search_word(driver,word)
    print (list2)
    for info in list2:
        if info !=None and info != '':
            with open('test.doc','w') as f:
                results = str(info)
                f.write(results)
                f.write("\n")
'''